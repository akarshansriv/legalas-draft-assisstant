from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re


def export_to_docx(text: str, filename: str = "petition.docx", key_dates: list = None) -> str:
    doc = Document()

    # Set page margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Define a consistent paragraph style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Parse the text and convert table sections to Word tables
    lines = text.strip().split("\n")
    i = 0
    processed_sections = set()  # Track processed table sections to avoid duplicates
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Check if this is a table section
        if _is_table_section(line) and line not in processed_sections:
            # Mark this section as processed
            processed_sections.add(line)
            # Add the section heading
            heading = doc.add_paragraph(line)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_format = heading.paragraph_format
            heading_format.space_after = Pt(12)
            
            # Make heading bold
            for run in heading.runs:
                run.bold = True
                run.font.size = Pt(14)
            
            # Skip empty lines after heading
            i += 1
            while i < len(lines) and lines[i].strip() == "":
                i += 1
            
            # Parse and create the table
            table_data = []
            while i < len(lines) and _is_table_row(lines[i]):
                row_data = _parse_table_row(lines[i])
                if row_data:
                    table_data.append(row_data)
                i += 1
            
            # Create Word table if we have data
            if table_data:
                _create_word_table(doc, table_data, line)
            else:
                # If no table data was parsed, create a default table structure
                _create_default_table(doc, line, key_dates)
            
            # Skip any remaining table-like content to prevent duplicates
            while i < len(lines) and (_is_table_row(lines[i]) or "|" in lines[i]):
                i += 1
            
            # Add some space after table
            doc.add_paragraph()
        else:
            # Regular paragraph
            if line == "":
                doc.add_paragraph()
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.first_line_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(8)
            i += 1

    # Save document
    os.makedirs("temp", exist_ok=True)
    output_path = os.path.join("temp", filename)
    doc.save(output_path)
    return output_path


def _is_table_section(line: str) -> bool:
    """Check if a line indicates the start of a table section."""
    table_sections = [
        "INDEX",
        "LIST OF DATED AND EVENTS"
    ]
    return any(section in line.upper() for section in table_sections)


def _is_table_row(line: str) -> bool:
    """Check if a line is a table row (contains | characters and is not a separator)."""
    if "|" not in line:
        return False
    
    # Check if it's a separator row (only dashes and spaces)
    stripped = line.replace("|", "").replace("-", "").replace(" ", "")
    if not stripped:
        return False
    
    # Check if it starts with a number (likely a data row)
    first_cell = line.split("|")[0].strip()
    if first_cell.isdigit() or first_cell in ["1", "2", "3", "4", "5", "6", "7", "8", "9"]:
        return True
    
    return False


def _parse_table_row(line: str) -> list:
    """Parse a table row and extract cell data."""
    # Remove leading/trailing whitespace and split by |
    cells = [cell.strip() for cell in line.split("|")]
    
    # Remove empty cells at the beginning and end
    while cells and cells[0] == "":
        cells.pop(0)
    while cells and cells[-1] == "":
        cells.pop()
    
    # Filter out separator rows (rows with only dashes and spaces)
    if all(cell.replace("-", "").replace(" ", "") == "" for cell in cells):
        return None
    
    return cells


def _create_word_table(doc: Document, table_data: list, section_name: str):
    """Create a Word table from the parsed data using the user's suggested approach."""
    if not table_data:
        return
    
    # Determine table structure based on section
    if "INDEX" in section_name.upper():
        headers = ["S. No.", "Particulars", "Page No."]
        col_widths = [0.8, 3.5, 1.0]
    elif "LIST OF DATED AND EVENTS" in section_name.upper():
        headers = ["S. No.", "Date", "Event Description"]
        col_widths = [0.8, 1.5, 3.0]
    else:
        # Generic table
        headers = table_data[0] if table_data else []
        col_widths = [1.0] * len(headers)
    
    # Create table with Word built-in style (no header row needed - section title is the header)
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = "Table Grid"  # This makes it look like a Word table with borders
    
    # Set column widths
    for i, width in enumerate(col_widths):
        if i < len(table.columns):
            table.columns[i].width = Inches(width)
    
    # Add data rows (include first header row, skip duplicate header rows)
    header_row_added = False
    for row_data in table_data:
        if len(row_data) >= len(headers):
            # Check if this is a header row
            first_cell = row_data[0].strip().lower() if row_data else ""
            is_header_row = any(header_word in first_cell for header_word in ["s.no", "s. no", "particulars", "description", "date", "page"])
            
            # If it's a header row and we haven't added one yet, add it
            if is_header_row and not header_row_added:
                header_row_added = True
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data[:len(headers)]):
                    if i < len(row_cells):
                        row_cells[i].text = cell_data
                
                # Format header cells
                for i, cell in enumerate(row_cells):
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.runs[0]
                        run.bold = True
                        run.font.size = Pt(11)
            # If it's a header row but we already added one, skip it (duplicate)
            elif is_header_row and header_row_added:
                continue
            # If it's a data row, add it
            else:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data[:len(headers)]):
                    if i < len(row_cells):
                        row_cells[i].text = cell_data
                
                # Formatting cells
                for i, cell in enumerate(row_cells):
                    for paragraph in cell.paragraphs:
                        if i == 1:  # Particulars/Description column -> left align
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            # Check if this is an annexure entry and make "ANNEXURE NO. X" bold
                            cell_text = cell.text
                            if "ANNEXURE NO." in cell_text.upper():
                                # Clear existing runs and create new ones with bold formatting for annexure
                                paragraph.clear()
                                run = paragraph.add_run()
                                run.font.size = Pt(11)
                                
                                # Split text to find annexure part
                                parts = cell_text.split("ANNEXURE NO.")
                                if len(parts) > 1:
                                    # Add text before "ANNEXURE NO."
                                    if parts[0]:
                                        run.text = parts[0]
                                    
                                    # Add "ANNEXURE NO." in bold
                                    annexure_part = "ANNEXURE NO." + parts[1]
                                    # Find where the description starts (after the colon)
                                    colon_index = annexure_part.find(":")
                                    if colon_index > 0:
                                        annexure_number = annexure_part[:colon_index + 1]
                                        description = annexure_part[colon_index + 1:]
                                        
                                        # Add annexure number in bold
                                        bold_run = paragraph.add_run(annexure_number)
                                        bold_run.bold = True
                                        bold_run.font.size = Pt(11)
                                        
                                        # Add description in normal text
                                        if description:
                                            normal_run = paragraph.add_run(description)
                                            normal_run.font.size = Pt(11)
                                    else:
                                        # If no colon found, make the whole annexure part bold
                                        bold_run = paragraph.add_run(annexure_part)
                                        bold_run.bold = True
                                        bold_run.font.size = Pt(11)
                                else:
                                    # Fallback: just set the text normally
                                    run.text = cell_text
                            else:
                                # Normal text formatting
                                run = paragraph.runs[0]
                                run.font.size = Pt(11)
                        else:  # S.No and Page No -> center align
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph.runs[0].font.size = Pt(11)
    
    # Add some space after the table
    doc.add_paragraph("\n")


def _create_default_table(doc: Document, section_name: str, key_dates: list = None):
    """Create a default table structure when no data is parsed."""
    if "INDEX" in section_name.upper():
        # Create default INDEX table (include header row)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        
        # Set column widths
        table.columns[0].width = Inches(0.8)
        table.columns[1].width = Inches(3.5)
        table.columns[2].width = Inches(1.0)
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "S. No."
        hdr_cells[1].text = "Particulars"
        hdr_cells[2].text = "Page No."
        
        # Format header
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0]
                run.bold = True
                run.font.size = Pt(11)
        
        # Add default rows
        default_rows = [
            ["1", "Index", "1"],
            ["2", "List of Dated and Events", "2"],
            ["3", "Application for Interim Relief", "3"],
            ["4", "Grounds", "4"],
            ["5", "Prayer", "5"],
            ["6", "Verification", "6"],
            ["7", "ANNEXURE NO. 1: Copy of the supporting document", "7"],
            ["8", "ANNEXURE NO. 2: Copy of the relevant order", "8"],
            ["9", "ANNEXURE NO. 3: Copy of the notice", "9"]
        ]
        
        for row_data in default_rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = cell_data
                # Format cells
                for paragraph in row_cells[i].paragraphs:
                    if i == 1:  # Particulars column -> left align
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Check if this is an annexure entry and make "ANNEXURE NO. X" bold
                        cell_text = cell_data
                        if "ANNEXURE NO." in cell_text.upper():
                            # Clear existing runs and create new ones with bold formatting for annexure
                            paragraph.clear()
                            
                            # Split text to find annexure part
                            parts = cell_text.split("ANNEXURE NO.")
                            if len(parts) > 1:
                                # Add text before "ANNEXURE NO."
                                if parts[0]:
                                    run = paragraph.add_run(parts[0])
                                    run.font.size = Pt(11)
                                
                                # Add "ANNEXURE NO." in bold
                                annexure_part = "ANNEXURE NO." + parts[1]
                                # Find where the description starts (after the colon)
                                colon_index = annexure_part.find(":")
                                if colon_index > 0:
                                    annexure_number = annexure_part[:colon_index + 1]
                                    description = annexure_part[colon_index + 1:]
                                    
                                    # Add annexure number in bold
                                    bold_run = paragraph.add_run(annexure_number)
                                    bold_run.bold = True
                                    bold_run.font.size = Pt(11)
                                    
                                    # Add description in normal text
                                    if description:
                                        normal_run = paragraph.add_run(description)
                                        normal_run.font.size = Pt(11)
                                else:
                                    # If no colon found, make the whole annexure part bold
                                    bold_run = paragraph.add_run(annexure_part)
                                    bold_run.bold = True
                                    bold_run.font.size = Pt(11)
                            else:
                                # Fallback: just set the text normally
                                run = paragraph.add_run(cell_text)
                                run.font.size = Pt(11)
                        else:
                            # Normal text formatting
                            paragraph.runs[0].font.size = Pt(11)
                    else:  # S.No and Page No -> center align
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.runs[0].font.size = Pt(11)
    
    elif "LIST OF DATED AND EVENTS" in section_name.upper():
        # Create default LIST OF DATED AND EVENTS table (include header row)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        
        # Set column widths
        table.columns[0].width = Inches(0.8)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(3.0)
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "S. No."
        hdr_cells[1].text = "Date"
        hdr_cells[2].text = "Event Description"
        
        # Format header
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0]
                run.bold = True
                run.font.size = Pt(11)
        
        # Add rows based on key_dates if available, otherwise use default
        if key_dates and len(key_dates) > 0:
            default_rows = []
            for i, date_event in enumerate(key_dates, 1):
                # Parse the date and event from the key_dates format
                if " - " in date_event:
                    date_part, event_part = date_event.split(" - ", 1)
                    default_rows.append([str(i), date_part.strip(), event_part.strip()])
                else:
                    # If no separator found, treat the whole string as event
                    default_rows.append([str(i), "Date not specified", date_event.strip()])
        else:
            # Fallback to default rows if no key_dates provided
            default_rows = [
                ["1", "DD.MM.YYYY", "Event description will be added here"],
                ["2", "DD.MM.YYYY", "Event description will be added here"],
                ["3", "DD.MM.YYYY", "Event description will be added here"]
            ]
        
        for row_data in default_rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = cell_data
                # Format cells
                for paragraph in row_cells[i].paragraphs:
                    if i == 2:  # Event Description column -> left align
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:  # S.No and Date -> center align
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.runs[0].font.size = Pt(11)
    
